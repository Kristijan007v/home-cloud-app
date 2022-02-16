from datetime import datetime
from operator import concat
from flask import *
from fpdf import FPDF
import ipinfo
from requests import get
from docx import Document
import zeep

time_now = datetime.now()
write_date = time_now.strftime('%a %b %d')
write_time = time_now.strftime('%a %b %d %Y %H:%M:%S ')

#defined current timestamp when saving files to insert into file name
file_wt = time_now.strftime('%d%m%Y-%H_%M_%S')


#Retrieve all the user data for logging purpose
def get_info():

    #Retrieve email from session and define file name for logs
    email = session['email']
    logName = f"{email}-login-logs.txt"

    #Get user IP
    get_ip = get('https://api.ipify.org').text
    ip_address = format(get_ip)

    #Get user location from IP using IPinfo API
    access_token = '156106a648a7eb'
    handler = ipinfo.getHandler(access_token)
    details = handler.getDetails(ip_address)
    location = details.city
    country = details.country

    #SOAP CLIENT to get full country name from iso country code provided py IPinfo API
    url = "http://webservices.oorsprong.org/websamples.countryinfo/CountryInfoService.wso?WSDL"
    client = zeep.Client(url)
    user_country = client.service.CountryName(country)

    #Get the browser user is using while visiting Home-cloud
    user_browser = request.headers.get('User-Agent')

    #Save retrieved info to log using log() function
    log_json(location, ip_address, user_browser, logName)
    log(location, ip_address, user_browser, user_country, logName)

    #Return ip and location for further use in application
    return ip_address, location



#Log all retrieved user data from get_info()
def log(location, user_ip, user_browser, user_country, logName):
    
    email = session ['email']
    f = open(f"static/Cloud/{email}/Logs/{logName}", "a")
    f.write("\n")
    f.write(write_time)
    f.write("\n")
    f.write(f"Country: {user_country}")
    f.write("\n")
    f.write(f"Location: {location}")
    f.write("\n")
    f.write(f"IP Address: {user_ip}")
    f.write("\n")
    f.write(f"Browser: {user_browser}")
    f.close()


#Convert user log file to pdf upon user request
def txt_to_pdf(logName):

    #Retrieve email from session and define file name for pdf
    email = session['email']
    filename = str(f"login-report-{file_wt}")

    # save FPDF() class into 
    # a variable pdf
    pdf = FPDF()   
    
    # Add a page
    pdf.add_page()
    
    # set style and size of font 
    # that you want in the pdf
    pdf.set_font("Arial", size = 13)
    
    # open the text file in read mode
    f = open(f"static/Cloud/{email}/Logs/{logName}", "r")
    
    # insert the texts in pdf
    for x in f:
        pdf.cell(200, 10, txt = x, ln = 1, align = 'L')
    
    # save the pdf with name .pdf
    pdf.output(f"static/Cloud/{email}/Folders/Reports/documents/{filename}.pdf")


def txt_to(extension, logName):
    
    email = session['email']
    txt_path = f"static/Cloud/{email}/Logs/{logName}"
    filename = str(f"login-report-{file_wt}.{extension}")
    save_file = f"static/Cloud/{email}/Folders/Reports/documents/{filename}"
    document = Document()
    document.add_heading(f'Login report - {email}', level=1)
    with open(txt_path) as txt:
        for line in txt:
            document.add_paragraph(line)
    document.save(save_file)


def save_json(filename, filesize, fileext, created_at, upload_path):
  
    # Data to be written
    dictionary ={
        "file_name" : str(f"{filename}"),
        "file_size" : str(f"{filesize}"),
        "file_exstension" : str(f"{fileext}"),
        "created_at" : str(f"{created_at}")
    }
    
    # Serializing json 
    json_object = json.dumps(dictionary, indent = 4)
    
    #Saving json file
    new_name = f"{filename}.json"
    save_to = f"{upload_path}/data/{new_name}"
    with open(save_to, "w") as outfile:
        outfile.write(json_object)


def log_json(location, user_ip, user_browser, logName):


    email = session['email']
    upload_path = f"static/Cloud/{email}/Logs"
    # Data to be written
    dictionary ={
        "time": str(f"{file_wt}"),
        "location" : str(f"{location}"),
        "user_ip" : str(f"{user_ip}"),
        "user_browser" : str(f"{user_browser}"),
    }
    
    # Serializing json 
    json_object = json.dumps(dictionary, indent = 4)
    
    #Saving json file
    new_name = f"{logName}.json"
    save_to = f"{upload_path}/{new_name}"
    with open(save_to, "a") as outfile:
        outfile.write(json_object)


#txt_to('rtf', 'kiki.vidovic.6969@gmail.com-login-logs.txt')

